import os
import sys
import json
import smtplib
import shutil
import traceback
import math
import time
from datetime import datetime, timedelta
from io import BytesIO
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.image import MIMEImage
from email.mime.application import MIMEApplication

import requests
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
from shapely.geometry import shape, Point
from shapely.ops import transform, nearest_points

# Configurar matplotlib para que funcione sin pantalla (servidor)
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import contextily as cx
from matplotlib.lines import Line2D

# Intentar importar librerías opcionales con seguridad
try:
    from pyproj import Transformer
except ImportError:
    Transformer = None
    print("⚠️ Advertencia: pyproj no está instalado. No se generarán mapas ni coordenadas GTM.", file=sys.stderr)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None
    print("⚠️ Advertencia: python-docx no está instalado. No se generará el informe Word.", file=sys.stderr)

MAP_KEY = "1f5837a949e2dff8572d9bb96df86898"

MESES_ES = {
    "01": "Enero", "02": "Febrero", "03": "Marzo", "04": "Abril",
    "05": "Mayo", "06": "Junio", "07": "Julio", "08": "Agosto",
    "09": "Septiembre", "10": "Octubre", "11": "Noviembre", "12": "Diciembre"
}

CAMPAMENTOS = [
    {"nombre": "Paxbán", "x": 541459.545, "y": 1968309.168},
    {"nombre": "El Guiro", "x": 546526.26, "y": 1966445.116},
    {"nombre": "Las corrientes", "x": 548522.604, "y": 1963905.742},
    {"nombre": "El Carrizal", "x": 549089.113, "y": 1960904.892},
    {"nombre": "Los campitos", "x": 542354.499, "y": 1962816.912},
    {"nombre": "Ojo de agua", "x": 539679.072, "y": 1960737.306},
    {"nombre": "El Tiempo", "x": 544849.236, "y": 1958376.547},
    {"nombre": "La Reforma", "x": 528989.788, "y": 1963524.435},
    {"nombre": "El Magueyal", "x": 526593.273, "y": 1967097.684},
    {"nombre": "El Mapache", "x": 520451.541, "y": 1965656.699},
    {"nombre": "La Pepesca", "x": 517819.874, "y": 1968240.009},
    {"nombre": "La Lagartija", "x": 517181.659, "y": 1969837.431},
    {"nombre": "El Salmoncito", "x": 517084.743, "y": 1963131.811},
    {"nombre": "El Salmon", "x": 517084.743, "y": 1962568.742},
    {"nombre": "Santa Elena", "x": 518382.504, "y": 1957356.105},
    {"nombre": "El Infiernon", "x": 521941.396, "y": 1953847.365},
    {"nombre": "La Isla", "x": 523325.068, "y": 1952239.194},
    {"nombre": "El Morgan", "x": 524265.663, "y": 1945609.637},
    {"nombre": "El Infiernito", "x": 529602.392, "y": 1954204.481},
    {"nombre": "El Jobal", "x": 536158.163, "y": 1966152.753},
    {"nombre": "Los Cuyos", "x": 533886.902, "y": 1957284.17},
    {"nombre": "El Recreo", "x": 537205.586, "y": 1965649.189},
    {"nombre": "Los Cuyos II", "x": 539968.134, "y": 1962862.671},
    {"nombre": "Los Perros", "x": 534691.126, "y": 1960808.64}
]

# Configuración oficial GTM (Guatemala Transversal Mercator)
GTM_PROJ_STR = "+proj=tmerc +lat_0=0 +lon_0=-90.5 +k=0.9998 +x_0=500000 +y_0=0 +ellps=WGS84 +datum=WGS84 +units=m +no_defs"

def convertir_a_gtm(lon, lat):
    """Convierte coordenadas de WGS84 (lat, lon) a GTM."""
    if not Transformer:
        return "No disponible"
    try:
        transformer = Transformer.from_crs(
            "EPSG:4326", 
            GTM_PROJ_STR,
            always_xy=True
        )
        gtm_x, gtm_y = transformer.transform(lon, lat)
        return f"{gtm_x:.2f} E, {gtm_y:.2f} N"
    except Exception as e:
        return "Error Calc"

def calcular_distancia_direccion(p, poly):
    """Calcula distancia y dirección desde un punto al polígono."""
    if not Transformer: return None
    try:
        # Proyección GTM para metros (EPSG:4326 -> GTM)
        trans_to_meter = Transformer.from_crs("EPSG:4326", GTM_PROJ_STR, always_xy=True)
        
        p_meter = transform(trans_to_meter.transform, p)
        poly_meter = transform(trans_to_meter.transform, poly)
        
        dist_meters = poly_meter.distance(p_meter)
        p_near = nearest_points(poly_meter, p_meter)[0]
        
        dx = p_meter.x - p_near.x
        dy = p_meter.y - p_near.y
        angle = math.degrees(math.atan2(dy, dx))
        if angle < 0: angle += 360

        if 337.5 <= angle or angle < 22.5: direction = "Este"
        elif 22.5 <= angle < 67.5: direction = "Noreste"
        elif 67.5 <= angle < 112.5: direction = "Norte"
        elif 112.5 <= angle < 157.5: direction = "Noroeste"
        elif 157.5 <= angle < 202.5: direction = "Oeste"
        elif 202.5 <= angle < 247.5: direction = "Suroeste"
        elif 247.5 <= angle < 292.5: direction = "Sur"
        else: # 292.5 <= angle < 337.5
            direction = "Sureste"
        
        return f"{int(dist_meters)} metros del límite {direction}"
    except Exception:
        return None

def calcular_campamento_cercano(lon, lat):
    """Calcula el campamento más cercano y la distancia."""
    if not Transformer: return "N/A"
    try:
        trans_to_gtm = Transformer.from_crs("EPSG:4326", GTM_PROJ_STR, always_xy=True)
        fx, fy = trans_to_gtm.transform(lon, lat)
        min_dist = float('inf')
        nearest_camp = None
        nearest_coords = None
        for c in CAMPAMENTOS:
            # Distancia Euclidiana
            dist = math.sqrt((fx - c['x'])**2 + (fy - c['y'])**2)
            if dist < min_dist:
                min_dist = dist
                nearest_camp = c['nombre']
                nearest_coords = (c['x'], c['y'])
        if nearest_camp:
            cx, cy = nearest_coords
            angle = math.degrees(math.atan2(fy - cy, fx - cx))
            if angle < 0: angle += 360
            
            if 337.5 <= angle or angle < 22.5: card = "al Este"
            elif 22.5 <= angle < 67.5: card = "al Noreste"
            elif 67.5 <= angle < 112.5: card = "al Norte"
            elif 112.5 <= angle < 157.5: card = "al Noroeste"
            elif 157.5 <= angle < 202.5: card = "al Oeste"
            elif 202.5 <= angle < 247.5: card = "al Suroeste"
            elif 247.5 <= angle < 292.5: card = "al Sur"
            else: card = "al Sureste"
            
            return f"{int(min_dist)}m de {nearest_camp} {card}"
    except Exception:
        pass
    return "N/A"

def enviar_correo_alerta(cuerpo_html, asunto="🔥 Alerta Paxban", imagen_mapa=None, archivo_zip=None):
    """Envía un correo electrónico de alerta."""
    SMTP_SERVER = os.environ.get("SMTP_SERVER")
    SMTP_PORT = os.environ.get("SMTP_PORT")
    SMTP_USER = os.environ.get("SMTP_USER")
    SMTP_PASSWORD = os.environ.get("SMTP_PASSWORD")
    RECIPIENT_EMAIL = os.environ.get("RECIPIENT_EMAIL")

    if not all([SMTP_SERVER, SMTP_PORT, SMTP_USER, SMTP_PASSWORD, RECIPIENT_EMAIL]):
        print("❌ Error: Faltan credenciales de correo. Revise los Secrets de GitHub.", file=sys.stderr)
        return

    destinatarios = [email.strip() for email in RECIPIENT_EMAIL.split(',') if email.strip()]
    fecha_hora = (datetime.utcnow() - timedelta(hours=6)).strftime("%d/%m/%Y %H:%M")
    asunto_completo = f"{asunto} - {fecha_hora}"

    print(f"📧 Enviando correo a: {', '.join(destinatarios)}")
    
    try:
        msg = MIMEMultipart()
        msg['From'] = SMTP_USER
        msg['To'] = ", ".join(destinatarios)
        msg['Subject'] = asunto_completo
        
        msg.attach(MIMEText(cuerpo_html, 'html', 'utf-8'))
        
        # Adjuntar logo
        try:
            with open('logo (2).png', 'rb') as f:
                logo_img = MIMEImage(f.read())
                logo_img.add_header('Content-ID', '<logo_paxban>')
                msg.attach(logo_img)
        except FileNotFoundError:
            print("⚠️ Advertencia: No se encontró logo (2).png. El correo se enviará sin logo.", file=sys.stderr)

        if imagen_mapa:
            img = MIMEImage(imagen_mapa)
            img.add_header('Content-ID', '<mapa_peten>')
            img.add_header('Content-Disposition', 'inline', filename='mapa_peten.png')
            msg.attach(img)
            
        if archivo_zip:
            nombre_archivo, datos_archivo = archivo_zip
            part = MIMEApplication(datos_archivo, Name=nombre_archivo)
            part['Content-Disposition'] = f'attachment; filename="{nombre_archivo}"'
            msg.attach(part)

        with smtplib.SMTP(SMTP_SERVER, int(SMTP_PORT)) as server:
            server.starttls()
            server.login(SMTP_USER, SMTP_PASSWORD)
            server.send_message(msg)
        print("✅ Correo enviado exitosamente.")
    except Exception as e:
        print(f"❌ Error enviando correo: {e}", file=sys.stderr)

def enviar_alerta_telegram(mensaje, imagen_bytes=None):
    """Envía mensaje a Telegram."""
    BOT_TOKEN = os.environ.get("TELEGRAM_BOT_TOKEN")
    CHAT_IDS_RAW = os.environ.get("TELEGRAM_CHAT_ID")

    if not all([BOT_TOKEN, CHAT_IDS_RAW]):
        return
    
    chat_ids = [cid.strip() for cid in CHAT_IDS_RAW.split(',') if cid.strip()]

    for chat_id in chat_ids:
        api_url = f"https://api.telegram.org/bot{BOT_TOKEN}/"
        data = {'chat_id': chat_id, 'parse_mode': 'HTML'}
        files = None

        if imagen_bytes:
            api_url += "sendPhoto"
            files = {'photo': ('mapa.png', imagen_bytes, 'image/png')}
            data['caption'] = mensaje
        else:
            api_url += "sendMessage"
            data['text'] = mensaje

        try:
            if files:
                requests.post(api_url, data=data, files=files, timeout=20)
            else:
                requests.post(api_url, json=data, timeout=10)
        except Exception as e:
            print(f"⚠️ Error Telegram ({chat_id}): {e}", file=sys.stderr)

def guardar_mapa_local(imagen_bytes):
    """Guarda el mapa en la carpeta mapa_reporte_diario."""
    if not imagen_bytes: return
    fecha_dt = datetime.utcnow() - timedelta(hours=6)
    mes_nombre = MESES_ES.get(fecha_dt.strftime("%m"), fecha_dt.strftime("%m"))
    carpeta = os.path.join("mapa_reporte_diario", fecha_dt.strftime("%Y"), mes_nombre)
    os.makedirs(carpeta, exist_ok=True)
    ruta = os.path.join(carpeta, f"Mapa_Calor_{fecha_dt.strftime('%Y-%m-%d')}.png")
    with open(ruta, "wb") as f: f.write(imagen_bytes)
    print(f"💾 Mapa guardado: {ruta}")

def guardar_bitacora(imagen_bytes, tipo, datos_puntos):
    """Guarda evidencia en bitácora."""
    if not imagen_bytes: return
    fecha_dt = datetime.utcnow() - timedelta(hours=6)
    carpeta = os.path.join("bitacora", fecha_dt.strftime("%Y"), fecha_dt.strftime("%m"), tipo)
    os.makedirs(carpeta, exist_ok=True)
    fecha_str = fecha_dt.strftime("%Y-%m-%d_%H%M")
    with open(os.path.join(carpeta, f"{tipo}_{fecha_str}.png"), "wb") as f: f.write(imagen_bytes)
    with open(os.path.join(carpeta, f"{tipo}_{fecha_str}.json"), "w", encoding="utf-8") as f:
        json.dump(datos_puntos, f, indent=2)

def limpiar_descargas_antiguas():
    """Elimina carpetas de meses con formato numérico si existe el formato nombre."""
    if not os.path.exists("descargas"): return
    
    print("🧹 Limpiando carpetas duplicadas en descargas...")
    for anio in os.listdir("descargas"):
        anio_path = os.path.join("descargas", anio)
        if not os.path.isdir(anio_path): continue
        
        for item in os.listdir(anio_path):
            # Si es carpeta numérica (ej: "1") y existe su par ("Enero"), borrar la numérica
            if item.isdigit():
                item_path = os.path.join(anio_path, item)
                if os.path.isdir(item_path):
                    mes_num = item.zfill(2)
                    mes_nombre = MESES_ES.get(mes_num)
                    if mes_nombre and os.path.exists(os.path.join(anio_path, mes_nombre)):
                        print(f"🗑️ Borrando duplicado obsoleto: {item_path}")
                        shutil.rmtree(item_path, ignore_errors=True)

def descargar_puntos_historicos(fecha_inicio, fecha_fin):
    """Descarga TODOS los puntos de calor de la región para un rango de fechas."""
    puntos = []
    session = requests.Session()
    retry = Retry(total=3, backoff_factor=1)
    adapter = HTTPAdapter(max_retries=retry)
    session.mount("https://", adapter)
    
    start = datetime.strptime(fecha_inicio, "%Y-%m-%d").date()
    end = datetime.strptime(fecha_fin, "%Y-%m-%d").date()
    delta = end - start
    dias = [start + timedelta(days=i) for i in range(delta.days + 1)]
    
    satelites = ["MODIS_NRT", "VIIRS_SNPP_NRT", "VIIRS_NOAA20_NRT"]
    headers = {"User-Agent": "PaxbanBot/1.0"}
    
    print(f"📡 Descargando historial regional del {fecha_inicio} al {fecha_fin}...")
    
    hoy = datetime.utcnow().date()
    
    for dia in dias:
        if dia > hoy:
            continue # Evitar descargar datos de días futuros que la API podría rellenar con datos actuales
            
        fecha_str = dia.strftime("%Y-%m-%d")
        for sat in satelites:
            try:
                # API para 1 día específico en el área
                url = f"https://firms.modaps.eosdis.nasa.gov/api/area/csv/{MAP_KEY}/{sat}/-94,13.5,-88,20/1/{fecha_str}"
                res = session.get(url, headers=headers, timeout=10)
                if res.status_code == 200:
                    lines = res.text.strip().split('\n')[1:]
                    for line in lines:
                        try:
                            # Limpiar espacios y comillas de cada campo para evitar errores de comparación
                            d = [x.strip().replace('"', '') for x in line.split(',')]
                            
                            # --- FILTRO DE SEGURIDAD DE FECHA ---
                            # Usamos objetos de fecha para evitar errores de formato (ej: 2026-2-6 vs 2026-02-06)
                            if len(d) > 5:
                                try:
                                    fecha_dato = datetime.strptime(d[5], "%Y-%m-%d").date()
                                    fecha_req = datetime.strptime(fecha_str, "%Y-%m-%d").date()
                                    if fecha_dato != fecha_req:
                                        continue # Ignorar si la fecha no coincide exactamente
                                except: pass
                            
                            # Convertir a hora Guatemala (UTC-6)
                            dt = datetime.strptime(f"{d[5]} {d[6]}", "%Y-%m-%d %H%M")
                            dt_gt = dt - timedelta(hours=6)
                            fecha_gt = dt_gt.strftime("%d/%m/%Y %H:%M")

                            puntos.append({
                                "lat": float(d[0]), "lon": float(d[1]), 
                                "color": "red", # Color base para mapas semanales
                                "sat": sat,
                                "fecha": f"{fecha_gt} (Hora GT)",
                                "fecha_simple": d[5] # Para filtrado fácil
                            })
                        except: pass
            except requests.exceptions.ConnectionError:
                print("⚠️ Error de conexión detectado. Deteniendo descarga histórica para evitar bloqueos.")
                return puntos
            except: pass
            time.sleep(0.1) # Evitar saturación
    return puntos

def generar_galeria_html():
    """Genera reportes.html."""
    try:
        mapas = []
        # Escanear tanto la carpeta estándar como la carpeta '2.4' detectada
        for base_dir in ["mapa_reporte_diario", "2.4"]:
            if os.path.exists(base_dir):
                for root, _, files in os.walk(base_dir):
                    for file in files:
                        if file.endswith(".png"):
                            url = os.path.join(root, file).replace(os.sep, '/')
                            mapas.append({"url": url, "nombre": file, "fecha": file.replace("Mapa_Calor_", "").replace(".png", "")})
        
        mapas.sort(key=lambda x: x['fecha'], reverse=True)
        
        # Buscar reportes mensuales (ZIPs)
        reportes_dict = {}
        if os.path.exists("descargas"):
            for root, _, files in os.walk("descargas"):
                for file in files:
                    if file.endswith(".zip"):
                        # --- FILTRO ANTI-DUPLICADOS ---
                        # Si el archivo está en una carpeta numérica (ej: descargas/2026/1)
                        # y existe la carpeta correcta (descargas/2026/Enero), lo ignoramos.
                        try:
                            folder_name = os.path.basename(root)
                            if folder_name.isdigit():
                                mes_nombre_check = MESES_ES.get(folder_name.zfill(2))
                                parent_folder = os.path.dirname(root)
                                if mes_nombre_check and os.path.exists(os.path.join(parent_folder, mes_nombre_check)):
                                    continue 
                        except: pass
                        # ------------------------------

                        url = os.path.join(root, file).replace(os.sep, '/')
                        
                        # Intentar crear un nombre bonito para la tarjeta (Ej: "Enero 2026")
                        nombre_mostrar = file
                        key = file
                        try:
                            # Formato esperado: Reporte_Mensual_01_2026.zip
                            parts = file.replace(".zip", "").split("_")
                            if len(parts) >= 4:
                                # Forzar conversión a entero para asegurar formato 01, 02...
                                mes_num = f"{int(parts[2]):02d}"
                                anio = parts[3]
                                mes_nombre = MESES_ES.get(mes_num, mes_num)
                                nombre_mostrar = f"{mes_nombre} {anio}"
                                key = f"{anio}-{mes_num}"
                        except: pass
                        
                        # Evitar duplicados por mes (preferir nombre de archivo más largo ej: 01 vs 1)
                        if key in reportes_dict:
                            print(f"🔹 Unificando reporte duplicado: {file} se agrupa bajo {key}")
                            
                        if key not in reportes_dict or len(file) > len(reportes_dict[key]['filename']):
                            reportes_dict[key] = {"url": url, "nombre": nombre_mostrar, "filename": file}
        
        reportes_mensuales = list(reportes_dict.values())
        print(f"📦 Se encontraron {len(reportes_mensuales)} reportes mensuales.")
        # Ordenar por nombre de archivo original para mantener orden cronológico
        reportes_mensuales.sort(key=lambda x: x['filename'], reverse=True)

        html = """<!DOCTYPE html><html lang="es"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0"><title>Reportes Paxban</title><link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
        <style>
            .card-img-top { height: 200px; object-fit: cover; }
            @media (max-width: 768px) { 
                h1 { font-size: 1.5rem; } 
                .container { padding-left: 15px; padding-right: 15px; }
                /* En móvil el botón es sólido para que no se vea como un bloque blanco vacío */
                .btn-pc-outline { background-color: #198754 !important; color: white !important; width: 100%; }
            }
            @media (min-width: 769px) {
                .btn-pc-outline { width: auto; }
            }
        </style>
        </head><body class="bg-light">
        <div class="container py-4 py-md-5">
            <div class="d-flex flex-column flex-md-row justify-content-between align-items-center mb-4 gap-3">
                <h1 class="text-success fw-bold m-0">📂 Galería de Reportes</h1>
                <a href="index.html" class="btn btn-outline-success btn-pc-outline px-4 shadow-sm">🏠 Volver al Inicio</a>
            </div>"""
        
        # Sección Reportes Mensuales
        if reportes_mensuales:
            html += """<h3 class="text-secondary mt-4 border-bottom pb-2">📦 Reportes Mensuales (Descarga Completa)</h3><div class="row row-cols-1 row-cols-md-3 g-4 mb-5">"""
            for r in reportes_mensuales:
                # Agregamos ?t=timestamp para evitar que el navegador use una versión vieja del ZIP en caché
                html += f"""<div class="col"><div class="card h-100 shadow-sm border-success"><div class="card-body text-center"><h5 class="card-title text-success">📅 {r['nombre']}</h5><p class="card-text small text-muted">Incluye: Reportes diarios, Incendios y Word.</p><a href="{r['url']}?t={int(time.time())}" class="btn btn-success w-100" download>⬇️ Descargar ZIP</a></div></div></div>"""
            html += "</div>"

        # Sección Mapas Diarios
        html += """<h3 class="text-secondary mt-4 border-bottom pb-2">🗺️ Mapas Diarios</h3><div class="row row-cols-1 row-cols-md-3 g-4">"""
        for m in mapas:
            html += f"""<div class="col"><div class="card h-100 shadow-sm"><img src="{m['url']}" class="card-img-top" style="height:250px;object-fit:cover;"><div class="card-body"><h5 class="card-title">{m['fecha']}</h5><a href="{m['url']}" class="btn btn-primary btn-sm" download target="_blank">⬇️ Descargar</a></div></div></div>"""
        html += "</div></div></body></html>"
        
        with open("reportes.html", "w", encoding="utf-8") as f: f.write(html)
        print("✅ Galería HTML generada.")
    except Exception as e:
        print(f"❌ Error generando galería: {e}", file=sys.stderr)

def generar_mapa_imagen(puntos, concesiones=None, center_point=None, buffer=0.1, basemap_provider=cx.providers.Esri.NatGeoWorldMap, is_pre_alert_map=False, draw_buffer=False):
    """Genera imagen PNG del mapa."""
    if not Transformer: return None
    try:
        transformer = Transformer.from_crs("EPSG:4326", "EPSG:3857", always_xy=True)
        xs, ys, colores = [], [], []

        # Buscar el polígono de Paxbán para usarlo después
        paxban_poly = None
        if concesiones:
            for nombre, poly in concesiones.items():
                if "Paxbán" in nombre:
                    paxban_poly = poly
                    break

        for p in puntos:
            x, y = transformer.transform(p['lon'], p['lat'])
            xs.append(x); ys.append(y); colores.append(p['color'])
        
        fig, ax = plt.subplots(figsize=(10, 10))
        
        if draw_buffer and paxban_poly:
            buffer_poly = paxban_poly.buffer(0.09) # Approx 10km
            buffer_3857 = transform(transformer.transform, buffer_poly)
            if buffer_3857.geom_type == 'Polygon':
                x_b, y_b = buffer_3857.exterior.xy
                ax.plot(x_b, y_b, color='orange', linestyle='--', linewidth=2, zorder=1)
            elif buffer_3857.geom_type == 'MultiPolygon':
                 for p_b in buffer_3857.geoms:
                    x_b, y_b = p_b.exterior.xy
                    ax.plot(x_b, y_b, color='orange', linestyle='--', linewidth=2, zorder=1)
        
        if concesiones:
            if paxban_poly:
                poly_3857 = transform(transformer.transform, paxban_poly)
                if poly_3857.geom_type == 'Polygon':
                    x, y = poly_3857.exterior.xy
                    ax.plot(x, y, color='#2e7d32', linewidth=2, zorder=1)
                elif poly_3857.geom_type == 'MultiPolygon':
                    for p in poly_3857.geoms:
                        x, y = p.exterior.xy
                        ax.plot(x, y, color='#2e7d32', linewidth=2, zorder=1)

        if xs: ax.scatter(xs, ys, c=colores, s=50, edgecolors='white', zorder=2)

        if is_pre_alert_map and len(puntos) == 1 and paxban_poly:
            punto_prealerta = puntos[0]
            p_geom = Point(punto_prealerta['lon'], punto_prealerta['lat'])
            p_3857 = transform(transformer.transform, p_geom)
            poly_3857 = transform(transformer.transform, paxban_poly)

            p_on_poly_3857, _ = nearest_points(poly_3857, p_3857)

            line_x = [p_on_poly_3857.x, p_3857.x]
            line_y = [p_on_poly_3857.y, p_3857.y]
            ax.plot(line_x, line_y, color='cyan', linestyle='--', linewidth=2, zorder=3)

            dist_text = punto_prealerta.get('dist_info', '')
            if 'metros' in dist_text:
                text_x = (p_on_poly_3857.x + p_3857.x) / 2
                text_y = (p_on_poly_3857.y + p_3857.y) / 2
                ax.text(text_x, text_y, f" {dist_text.split(' ')[0]}m ", color='white', fontsize=14,
                        ha='center', va='center', zorder=4,
                        bbox=dict(facecolor='black', alpha=0.6, boxstyle='round,pad=0.2', edgecolor='cyan'))

        if center_point:
            lon, lat = center_point
            minx, miny = transformer.transform(lon - buffer, lat - buffer)
            maxx, maxy = transformer.transform(lon + buffer, lat + buffer)
        elif draw_buffer and paxban_poly:
            poly_3857 = transform(transformer.transform, paxban_poly)
            bounds = poly_3857.buffer(0.1).bounds # Usar un buffer un poco más grande para el view
            minx, miny, maxx, maxy = bounds
        else:
            minx, miny = transformer.transform(-91.5, 15.8)
            maxx, maxy = transformer.transform(-89.0, 17.9)

        ax.set_xlim(minx, maxx); ax.set_ylim(miny, maxy)
        
        try:
            cx.add_basemap(ax, crs="EPSG:3857", source=basemap_provider, attribution=False)
        except Exception as e:
            print(f"⚠️ Advertencia: No se pudo descargar el mapa base ({e}). Se generará sin fondo.", file=sys.stderr)
            
        ax.set_axis_off()

        if draw_buffer:
            buffer_line = Line2D([0], [0], color='orange', lw=2, linestyle='--', label='Zona de Amortiguamiento (10km)')
            paxban_line = Line2D([0], [0], color='#2e7d32', lw=2, label='Límite Concesión Paxbán')
            ax.legend(handles=[paxban_line, buffer_line], loc='upper right', frameon=True, facecolor='white', framealpha=0.8)
        
        buf = BytesIO()
        plt.savefig(buf, format='png', bbox_inches='tight', pad_inches=0)
        buf.seek(0); plt.close(fig)
        return buf.read()
    except Exception as e:
        print(f"⚠️ Error generando mapa: {e}", file=sys.stderr)
        return None

def crear_informe_word(ruta_salida, mes_nombre, anio, fires_list, map_images, concesiones=None, pre_alerts_list=None):
    if not Document: return
    try:
        doc = Document()
        
        # Configuración de estilo base (Arial 11)
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        # --- ENCABEZADO ---
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = "SISTEMA DE ALERTA TEMPRANA PAXBAN"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_para.style.font.bold = True
        header_para.style.font.size = Pt(14)
        header_para.style.font.color.rgb = RGBColor(0, 0, 0) # Negro Formal

        # --- LOGO (Derecha) ---
        try:
            logo_file = 'logo-giborv2 (2) (1).png'
            if not os.path.exists(logo_file): logo_file = 'logo (2).png' # Respaldo
            if os.path.exists(logo_file):
                logo_para = doc.add_paragraph()
                logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = logo_para.add_run()
                run.add_picture(logo_file, width=Inches(1.5))
        except Exception: pass

        # --- TÍTULOS CENTRADOS ---
        p_titles = doc.add_paragraph()
        p_titles.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_titles.add_run("GIBOR, S.A.\n")
        run.bold = True
        run.font.size = Pt(16)
        
        run = p_titles.add_run("UNIDAD DE MANEJO PAXBAN\n")
        run.bold = True
        run.font.size = Pt(14)
        
        run = p_titles.add_run(f"\nINFORME MENSUAL: {mes_nombre.upper()} {anio}")
        run.bold = True
        run.font.size = Pt(12)

        doc.add_paragraph() # Espacio

        # --- CONTEXTO NORMATIVO ---
        p_context = doc.add_paragraph()
        p_context.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p_context.add_run("CONTEXTO NORMATIVO Y OPERATIVO\n")
        run.bold = True
        p_context.add_run(
            "Este informe responde al Programa de Planificación, Monitoreo y Evaluación, principalmente al inciso 2.4: "
            "\"Protocolo para el seguimiento de punto de calor en la zona\", durante la temporada crítica. "
            "El sistema garantiza una vigilancia ininterrumpida sobre el patrimonio natural bajo custodia."
        )

        # --- UBICACIÓN Y ALCANCE ---
        p_loc = doc.add_paragraph()
        p_loc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p_loc.add_run("ALCANCE GEOGRÁFICO\n")
        run.bold = True
        p_loc.add_run(
            "Área Monitoreada: 65,755 hectáreas (todo el polígono de Paxbán) las 24 horas.\n"
            "Ubicación: Reserva de la Biosfera Maya, Zona de Usos Múltiples, Petén, Guatemala."
        )

        # --- INTRODUCCIÓN TÉCNICA ---
        doc.add_heading('1. Introducción Técnica', level=1)
        intro_text = (
            "El presente documento constituye el informe técnico mensual generado por el Sistema de Alerta Temprana Paxbán, "
            "una herramienta tecnológica de vanguardia implementada para la vigilancia permanente y detección oportuna de "
            "anomalías térmicas. Este software opera de manera automatizada en la nube, integrando datos satelitales de "
            "alta resolución en tiempo real con análisis geoespacial preciso.\n\n"
            "El objetivo primordial es fortalecer las capacidades de respuesta rápida y proporcionar información crítica "
            "para la gestión forestal sostenible y la protección de la biodiversidad, en estricto cumplimiento con los "
            "lineamientos técnicos del Consejo Nacional de Áreas Protegidas (CONAP) y los estándares de certificación FSC."
        )
        doc.add_paragraph(intro_text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # --- METODOLOGÍA ---
        doc.add_heading('2. Metodología y Tecnología', level=1)
        metodo_text = (
            "La arquitectura del sistema se basa en la integración de múltiples capas de tecnología para garantizar la precisión:\n\n"
            "A. Sensores Satelitales: Se realiza un monitoreo constante mediante los sensores MODIS (Moderate Resolution Imaging "
            "Spectroradiometer) y VIIRS (Visible Infrared Imaging Radiometer Suite) a través de la plataforma NASA FIRMS. "
            "Estos sensores permiten la detección de anomalías térmicas con una resolución espacial de hasta 375 metros.\n\n"
            "B. Procesamiento Geoespacial: El sistema ejecuta algoritmos de intersección espacial utilizando librerías avanzadas "
            "(Shapely, Pyproj). Esto permite discriminar automáticamente entre incendios dentro de la concesión, en la zona de "
            "amortiguamiento (buffer de 10 km) y eventos externos.\n\n"
            "C. Proyección Cartográfica: Para facilitar la navegación de las cuadrillas de campo, todas las coordenadas geográficas (WGS84) "
            "son transformadas automáticamente al sistema GTM (Guatemala Transversal Mercator), estándar oficial para la cartografía nacional."
        )
        doc.add_paragraph(metodo_text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # --- RESULTADOS ---
        doc.add_heading('3. Resultados del Monitoreo Mensual', level=1)

        if not fires_list and not pre_alerts_list:
            resumen_p = doc.add_paragraph(
                f"Durante el periodo correspondiente al mes de {mes_nombre}, el sistema mantuvo un monitoreo ininterrumpido sobre la Unidad de Manejo Paxbán. "
                "Tras el análisis exhaustivo de los datos satelitales, se concluye que el periodo finalizó SIN NINGUNA ALERTA de incendio "
                "dentro de la concesión y SIN PRE-ALERTAS en la zona de amortiguamiento.\n\n"
                "Este resultado representa un escenario de estabilidad ideal para la cobertura forestal e indica la efectividad de las medidas de prevención implementadas."
            )
            resumen_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            resumen_p = doc.add_paragraph(
                f"Durante el periodo correspondiente al mes de {mes_nombre}, el sistema de monitoreo satelital registró la siguiente actividad térmica en el área de interés y su zona de amortiguamiento. A continuación se presenta el desglose:"
            )
            resumen_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.add_heading('3.1. Alertas de Incendio (Interior de la Concesión)', level=2)
        if fires_list:
            table = doc.add_table(rows=0, cols=2)
            table.autofit = False
            table.allow_autofit = False
            table.columns[0].width = Inches(3.0)
            table.columns[1].width = Inches(3.0)

            for i in range(0, len(fires_list), 2):
                row_cells = table.add_row().cells
                
                # --- Evento 1 ---
                fire1 = fires_list[i]
                cell1 = row_cells[0]
                p1 = cell1.paragraphs[0] if cell1.paragraphs else cell1.add_paragraph()
                p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p1.paragraph_format.keep_with_next = True
                run1 = p1.add_run(f"Evento No. {i+1} - {fire1.get('fecha', 'N/A')}\n")
                run1.bold = True
                p1.add_run(f"Satélite: {fire1.get('sat', 'Desconocido')}\n")
                p1.add_run(f"Ubicación Referencial: {fire1.get('dist_campamento', 'N/A')}\n")
                p1.add_run(f"Coordenadas GTM: {fire1.get('gtm', 'N/A')}\n")
                
                try:
                    img_focused1 = generar_mapa_imagen([fire1], concesiones, center_point=(fire1['lon'], fire1['lat']), buffer=0.08, basemap_provider=cx.providers.Esri.WorldImagery)
                    if img_focused1:
                        pic_para1 = cell1.add_paragraph()
                        pic_para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run_pic1 = pic_para1.add_run()
                        run_pic1.add_picture(BytesIO(img_focused1), width=Inches(2.8))
                except: pass

                # --- Evento 2 (si existe) ---
                if i + 1 < len(fires_list):
                    fire2 = fires_list[i+1]
                    cell2 = row_cells[1]
                    p2 = cell2.paragraphs[0] if cell2.paragraphs else cell2.add_paragraph()
                    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p2.paragraph_format.keep_with_next = True
                    run2 = p2.add_run(f"Evento No. {i+2} - {fire2.get('fecha', 'N/A')}\n")
                    run2.bold = True
                    p2.add_run(f"Satélite: {fire2.get('sat', 'Desconocido')}\n")
                    p2.add_run(f"Ubicación Referencial: {fire2.get('dist_campamento', 'N/A')}\n")
                    p2.add_run(f"Coordenadas GTM: {fire2.get('gtm', 'N/A')}\n")

                    try:
                        img_focused2 = generar_mapa_imagen([fire2], concesiones, center_point=(fire2['lon'], fire2['lat']), buffer=0.08, basemap_provider=cx.providers.Esri.WorldImagery)
                        if img_focused2:
                            pic_para2 = cell2.add_paragraph()
                            pic_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run_pic2 = pic_para2.add_run()
                            run_pic2.add_picture(BytesIO(img_focused2), width=Inches(2.8))
                    except: pass
        else:
            p = doc.add_paragraph("No se registraron alertas de incendio dentro de los límites de la Concesión Industrial Paxbán durante este periodo.")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # --- ZONA DE AMORTIGUAMIENTO ---
        doc.add_heading('3.2. Zona de Amortiguamiento (Buffer de 10 km)', level=2)
        p_buffer_desc = doc.add_paragraph()
        p_buffer_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_buffer_desc.add_run(
            "Para la detección proactiva de amenazas, el sistema monitorea una zona de amortiguamiento que se extiende 10 kilómetros "
            "alrededor del perímetro de la Concesión Paxbán. El siguiente mapa ilustra esta área de vigilancia estratégica, "
            "donde cualquier punto de calor detectado se clasifica como una "
        )
        p_buffer_desc.add_run("pre-alerta.").italic = True

        try:
            buffer_map_bytes = generar_mapa_imagen([], concesiones, draw_buffer=True, basemap_provider=cx.providers.Esri.WorldImagery)
            if buffer_map_bytes:
                p_map = doc.add_paragraph()
                p_map.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_map = p_map.add_run()
                run_map.add_picture(BytesIO(buffer_map_bytes), width=Inches(5.0))
        except Exception as e:
            print(f"⚠️ Error generando mapa de zona de amortiguamiento: {e}")

        if pre_alerts_list:
            doc.add_heading('3.3. Puntos de Pre-Alerta en Zona de Amortiguamiento', level=2)
            table_pre = doc.add_table(rows=0, cols=2)
            table_pre.autofit = False
            table_pre.allow_autofit = False
            table_pre.columns[0].width = Inches(3.0)
            table_pre.columns[1].width = Inches(3.0)

            for i in range(0, len(pre_alerts_list), 2):
                row_cells = table_pre.add_row().cells
                
                # --- Pre-Alerta 1 ---
                pre_alert1 = pre_alerts_list[i]
                cell1 = row_cells[0]
                p1 = cell1.paragraphs[0] if cell1.paragraphs else cell1.add_paragraph()
                p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p1.paragraph_format.keep_with_next = True
                run1 = p1.add_run(f"Pre-Alerta No. {i+1} - {pre_alert1.get('fecha', 'N/A')}\n")
                run1.bold = True
                p1.add_run(f"Satélite: {pre_alert1.get('sat', 'Desconocido')}\n")
                p1.add_run(f"Ubicación Referencial: {pre_alert1.get('dist_info', 'N/A')}\n")
                p1.add_run(f"Coordenadas GTM: {pre_alert1.get('gtm', 'N/A')}\n")

                try:
                    img_focused1 = generar_mapa_imagen([pre_alert1], concesiones, center_point=(pre_alert1['lon'], pre_alert1['lat']), buffer=0.08, basemap_provider=cx.providers.Esri.WorldImagery, is_pre_alert_map=True)
                    if img_focused1:
                        pic_para1 = cell1.add_paragraph()
                        pic_para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run_pic1 = pic_para1.add_run()
                        run_pic1.add_picture(BytesIO(img_focused1), width=Inches(2.8))
                except: pass

                # --- Pre-Alerta 2 (si existe) ---
                if i + 1 < len(pre_alerts_list):
                    pre_alert2 = pre_alerts_list[i+1]
                    cell2 = row_cells[1]
                    p2 = cell2.paragraphs[0] if cell2.paragraphs else cell2.add_paragraph()
                    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p2.paragraph_format.keep_with_next = True
                    run2 = p2.add_run(f"Pre-Alerta No. {i+2} - {pre_alert2.get('fecha', 'N/A')}\n")
                    run2.bold = True
                    p2.add_run(f"Satélite: {pre_alert2.get('sat', 'Desconocido')}\n")
                    p2.add_run(f"Ubicación Referencial: {pre_alert2.get('dist_info', 'N/A')}\n")
                    p2.add_run(f"Coordenadas GTM: {pre_alert2.get('gtm', 'N/A')}\n")

                    try:
                        img_focused2 = generar_mapa_imagen([pre_alert2], concesiones, center_point=(pre_alert2['lon'], pre_alert2['lat']), buffer=0.08, basemap_provider=cx.providers.Esri.WorldImagery, is_pre_alert_map=True)
                        if img_focused2:
                            pic_para2 = cell2.add_paragraph()
                            pic_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run_pic2 = pic_para2.add_run()
                            run_pic2.add_picture(BytesIO(img_focused2), width=Inches(2.8))
                    except: pass
        # --- ANEXO GRÁFICO ---
        if map_images:
            doc.add_page_break()
            doc.add_heading('4. Reportes de Monitoreo Semanal', level=1)
            doc.add_paragraph("A continuación se presentan los mapas de calor acumulados por semana, mostrando TODA la actividad térmica detectada en la región (incluyendo quemas agrícolas y eventos externos) para verificar la cobertura del monitoreo.").alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            for img_path in map_images:
                if os.path.exists(img_path):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(img_path, width=Inches(6.0))
                    nombre_limpio = os.path.basename(img_path).replace(".png", "").replace("Mapa_Semanal_", "").replace("_", " ")
                    caption = doc.add_paragraph(f"{nombre_limpio}\n(Puntos acumulados del periodo)")
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.style.font.size = Pt(9)
                    caption.style.font.italic = True

        # --- CONCLUSIONES ---
        doc.add_heading('5. Conclusiones y Recomendaciones', level=1)
        concl_text = (
            "El Sistema de Alerta Temprana ha demostrado ser una herramienta eficaz para la gestión de riesgos en la Unidad de Manejo Paxbán. "
            "La integración de datos en tiempo real permite reducir significativamente los tiempos de respuesta ante conatos de incendio.\n\n"
            "Se recomienda:\n"
            "1. Continuar con la verificación de campo de cualquier alerta generada, incluso aquellas en la zona de amortiguamiento.\n"
            "2. Mantener actualizados los equipos de comunicación para la recepción de las alertas vía Telegram y Correo Electrónico.\n"
            "3. Utilizar las coordenadas GTM proporcionadas en este informe para la planificación de rutas de acceso en caso de emergencia."
        )
        doc.add_paragraph(concl_text).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Firma Final
        doc.add_paragraph("\n\n\n")
        firma = doc.add_paragraph("__________________________\nSistema de Alerta Temprana Paxban")
        firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
        firma_run = firma.add_run("\nDesarrollado por\nNery Jose Corado Ramírez\nMiembro de la CIF\nGIBOR, S.A")
        firma_run.italic = True
        firma_run.font.size = Pt(10)

        doc.save(ruta_salida)
    except Exception as e:
        print(f"⚠️ Error creando Word: {e}", file=sys.stderr)

def generar_reporte_mensual(concesiones):
    """Genera ZIP mensual."""
    print("📦 Iniciando generación de Reporte Mensual...")
    try:
        fecha_dt = datetime.utcnow() - timedelta(hours=6)
        # Permitir especificar mes y año vía variables de entorno para reportes retroactivos
        anio = os.environ.get("TARGET_YEAR") or fecha_dt.strftime("%Y")
        mes = os.environ.get("TARGET_MONTH") or fecha_dt.strftime("%m")
        mes = mes.zfill(2) # Asegurar formato de 2 dígitos (01, 02...)
        nombre_mes = MESES_ES.get(mes, mes)
        
        raiz = f"Reporte_{mes}_{anio}"
        if os.path.exists(raiz): shutil.rmtree(raiz)
        os.makedirs(raiz)
        
        # Copiar carpetas
        for d in ["Reporte Diario", "Incendios Detectados", "Informe de Puntos de Calor"]:
            os.makedirs(os.path.join(raiz, d), exist_ok=True)
            
        # Copiar contenido de mapa_reporte_diario y también de '2.4' si existen para el mes
        fuentes = [
            os.path.join("mapa_reporte_diario", anio, nombre_mes),
            os.path.join("2.4", anio, mes) # Formato numérico
        ]
        
        for src in fuentes:
            if os.path.exists(src):
                print(f"📂 Copiando mapas desde: {src}")
                for f in os.listdir(src):
                    if f.endswith(".png"):
                        # --- FILTRO DE SEGURIDAD DE ARCHIVOS ---
                        # Solo copiar si el nombre del archivo coincide con el mes/año del reporte
                        # Esto evita que un mapa de Febrero (02) se cuele en el reporte de Enero (01)
                        if f"_{anio}-{mes}-" in f:
                            shutil.copy2(os.path.join(src, f), os.path.join(raiz, "Reporte Diario"))

        # --- RECONSTRUCCIÓN DE EVENTOS DESDE HISTORIAL NASA ---
        # En lugar de depender de las bitácoras (que pudieron no guardarse),
        # se descarga el historial completo del mes y se reclasifica cada punto.
        # Esto asegura que el reporte mensual sea siempre la fuente de verdad.
        print("🔎 Recalculando todos los eventos del mes desde el historial de NASA FIRMS para máxima precisión...")
        fires_details = []
        pre_alerts_details = []

        f_inicio_mes = f"{anio}-{mes}-01"
        ultimo_dia_mes = (datetime(int(anio), int(mes) % 12 + 1, 1) - timedelta(days=1)).day if int(mes) < 12 else 31
        f_fin_mes = f"{anio}-{mes}-{ultimo_dia_mes:02d}"

        puntos_del_mes = descargar_puntos_historicos(f_inicio_mes, f_fin_mes)

        for punto in puntos_del_mes:
            p = Point(punto['lon'], punto['lat'])
            en_paxban = False
            en_prealerta = False
            
            for nom, poly in concesiones.items():
                if "Paxbán" in nom:
                    if poly.contains(p):
                        en_paxban = True
                        punto['dist_campamento'] = calcular_campamento_cercano(punto['lon'], punto['lat'])
                        break
                    elif poly.distance(p) < 0.09: # Buffer de 10km
                        en_prealerta = True
                        punto['dist_info'] = calcular_distancia_direccion(p, poly)
            
            if en_paxban:
                punto['gtm'] = convertir_a_gtm(punto['lon'], punto['lat'])
                fires_details.append(punto)
            elif en_prealerta:
                punto['gtm'] = convertir_a_gtm(punto['lon'], punto['lat'])
                pre_alerts_details.append(punto)

        print(f"Total de eventos recalculados: {len(fires_details)} alertas, {len(pre_alerts_details)} pre-alertas.")

        # Copiar imágenes de la bitácora si existen (para evidencia)
        src_alertas_img = os.path.join("bitacora", anio, mes, "alertas")
        if os.path.exists(src_alertas_img):
            for f in os.listdir(src_alertas_img):
                if f.endswith(".png"):
                    shutil.copy2(os.path.join(src_alertas_img, f), os.path.join(raiz, "Incendios Detectados"))

        # --- GENERACIÓN DE MAPAS SEMANALES ACUMULADOS PARA EL WORD ---
        # Genera 4 mapas (cubriendo todo el mes) con todos los puntos de calor acumulados en ese periodo
        map_images_paths = []
        temp_maps_dir = os.path.join(raiz, "Mapas_Semanales_Word")
        os.makedirs(temp_maps_dir, exist_ok=True)
        
        periodos = [(1, 7), (8, 14), (15, 21), (22, 31)]
        
        for i, (inicio, fin) in enumerate(periodos):
            # Calcular fechas reales para la descarga
            try:
                # Ajustar fin si el mes tiene menos días
                ultimo_dia_mes = (datetime(int(anio), int(mes) % 12 + 1, 1) - timedelta(days=1)).day if int(mes) < 12 else 31
                fin_real = min(fin, ultimo_dia_mes)
                
                f_inicio = f"{anio}-{mes}-{inicio:02d}"
                f_fin = f"{anio}-{mes}-{fin_real:02d}"
                
                # Descargar DATOS REALES de la región para esa semana
                puntos_semana = [p for p in puntos_del_mes if f_inicio <= p['fecha_simple'] <= f_fin]
            
                print(f"🗺️ Generando Reporte de Monitoreo Semana {i+1} ({len(puntos_semana)} puntos regionales)...")
                img_bytes = generar_mapa_imagen(puntos_semana, concesiones)
                
                if img_bytes:
                    nombre_archivo = f"Mapa_Semanal_Reporte_Monitoreo_Semana_{i+1}_(Del_{inicio}_al_{fin_real}).png"
                    ruta_archivo = os.path.join(temp_maps_dir, nombre_archivo)
                    with open(ruta_archivo, "wb") as f:
                        f.write(img_bytes)
                    map_images_paths.append(ruta_archivo)
            except Exception as e:
                print(f"⚠️ Error procesando semana {i+1}: {e}")

        # Generar Word con el nuevo formato
        crear_informe_word(
            os.path.join(raiz, "Informe de Puntos de Calor", "Informe.docx"), 
            nombre_mes, anio, fires_details, map_images_paths,
            concesiones=concesiones, pre_alerts_list=pre_alerts_details
        )

        # ZIP
        zip_filename = f"Reporte_Mensual_{mes}_{anio}"
        shutil.make_archive(zip_filename, 'zip', root_dir='.', base_dir=raiz)
        
        # Mover a descargas organizado por Año/Mes (Nombre)
        carpeta_destino = os.path.join("descargas", anio, nombre_mes)
        os.makedirs(carpeta_destino, exist_ok=True)
        ruta_final = os.path.join(carpeta_destino, f"{zip_filename}.zip")
        if os.path.exists(ruta_final): os.remove(ruta_final) # Evitar error si existe
        shutil.move(f"{zip_filename}.zip", ruta_final)
        shutil.rmtree(raiz)
        
        print(f"✅ ZIP creado: {ruta_final}")
        
        with open(ruta_final, "rb") as f: zip_bytes = f.read()
        
        cuerpo = f"""
        <html><head><style>
            @media print {{
                @page {{ margin: 0.5cm; }} 
                body {{ font-family: Arial, sans-serif; font-size: 9pt; }} 
                h2 {{ color: #1565C0; margin-top: 0; font-size: 12pt; margin-bottom: 5px; }}
                .info-box {{ background-color: #e3f2fd !important; border-left: 5px solid #1565C0 !important; -webkit-print-color-adjust: exact; padding: 5px !important; margin: 5px 0 !important; }}
                table {{ margin-bottom: 5px !important; }}
                td {{ padding-bottom: 0 !important; }}
                .no-print {{ display: none; }}
            }}
        </style></head><body>
        <div style="font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto;">
            <table style="width: 100%; border-bottom: 2px solid #1565C0; margin-bottom: 15px;">
                <tr>
                    <td style="width: 100px; padding-bottom: 10px;">
                        <img src="cid:logo_paxban" alt="Logo Paxban" style="width: 90px; height: auto;">
                    </td>
                    <td style="vertical-align: middle; padding-bottom: 10px;">
                        <h2 style="color: #1565C0; margin: 0;">📦 Reporte Mensual Generado: {nombre_mes} {anio}</h2>
                    </td>
                </tr>
            </table>
            <p>Estimado usuario,</p>
            <p>Se ha completado la compilación del reporte mensual de monitoreo satelital.</p>

            <div class="info-box" style="background-color: #e3f2fd; padding: 15px; border-left: 5px solid #1565C0; margin: 20px 0;">
                <h3 style="margin: 0; color: #0d47a1;">Archivo Adjunto: {zip_filename}.zip</h3>
                <p style="margin: 5px 0 0 0;">El archivo ZIP adjunto contiene las siguientes carpetas:</p>
                <ul style="margin-top: 10px; padding-left: 20px;">
                    <li><strong>Reporte Diario:</strong> Todos los mapas de calor diarios del mes.</li>
                    <li><strong>Incendios Detectados:</strong> Evidencia de alertas de incendio (si las hubo).</li>
                    <li><strong>Informe de Puntos de Calor:</strong> Documento Word con el resumen.</li>
                </ul>
            </div>

            <p>Puede descargar el archivo directamente desde este correo.</p>
            <br><hr style="border: 0; border-top: 1px solid #eee;">
            <div style="font-size: 12px; color: #666;">
                <p><b>Sistema de Alerta Temprana Paxban</b><br>Mensaje generado automáticamente.<br>Desarrollado por JR23CR</p>
                <p style="text-align: center;" class="no-print"><a href="https://JR23CR.github.io/alerta-paxban/reportes.html" style="background-color: #1565C0; color: white; padding: 10px 20px; text-decoration: none; border-radius: 5px; font-weight: bold;">📂 Ir a la Galería de Reportes</a></p>
            </div>
        </div>
        </body></html>"""
        enviar_correo_alerta(cuerpo, asunto=f"Reporte Mensual {nombre_mes} {anio}", archivo_zip=(f"{zip_filename}.zip", zip_bytes))
        
    except Exception as e:
        print(f"❌ Error CRÍTICO en reporte mensual: {e}", file=sys.stderr)
        traceback.print_exc()

def cargar_concesiones():
    try:
        with open('concesiones1.geojson', 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        if not Transformer:
            print("⚠️ Pyproj no disponible, no se puede aplicar desplazamiento al polígono.")
            return {f['properties'].get('Name', 'X'): shape(f['geometry']) for f in data['features']}

        concesiones_dict = {}
        trans_to_gtm = Transformer.from_crs("EPSG:4326", GTM_PROJ_STR, always_xy=True)
        trans_to_wgs = Transformer.from_crs(GTM_PROJ_STR, "EPSG:4326", always_xy=True)

        for f in data['features']:
            poly = shape(f['geometry'])
            name = f['properties'].get('Name', 'X')
            
            if "Paxbán" in name:
                print("🏗️ Aplicando desplazamiento de 100m al Norte al polígono de Paxbán para todos los cálculos...")
                poly_gtm = transform(trans_to_gtm.transform, poly)
                def shift_north(x, y, z=None):
                    return x, y + 100
                poly_shifted_gtm = transform(shift_north, poly_gtm)
                poly = transform(trans_to_wgs.transform, poly_shifted_gtm)

            concesiones_dict[name] = poly
        return concesiones_dict
    except Exception as e:
        print(f"⚠️ Error cargando concesiones: {e}", file=sys.stderr)
        return {}

def main():
    print("🚀 Iniciando sistema Paxban...")
    concesiones = cargar_concesiones()
    action_type = os.environ.get("ACTION_TYPE", "monitor")
    
    puntos = []
    
    # --- LÓGICA DE PRUEBAS (SIMULACROS) ---
    if action_type.startswith("test_"):
        print(f"🧪 MODO PRUEBA ACTIVADO: {action_type}")
        fecha_sim = (datetime.utcnow() - timedelta(hours=6)).strftime("%d/%m/%Y %H:%M")
        
        if action_type == "test_incendio":
            # Punto DENTRO de Paxbán
            puntos.append({
                "lat": 17.7, "lon": -90.15, "color": "red", "alerta": True, "pre_alerta": False,
                "sat": "SIMULACRO", "fecha": fecha_sim, "horas": 1,
                "concesion": "Paxbán", "gtm": convertir_a_gtm(-90.15, 17.7),
                "dist_info": None,
                "dist_campamento": calcular_campamento_cercano(-90.15, 17.7)
            })
        elif action_type == "test_prealerta":
            # Punto CERCA de Paxbán (Zona de Amortiguamiento)
            lat, lon = 17.55, -90.2
            p_test = Point(lon, lat)
            dist_info_test = "Zona de Amortiguamiento"
            
            # Calcular dinámicamente usando el polígono real
            for nom, poly in concesiones.items():
                if "Paxbán" in nom:
                    res = calcular_distancia_direccion(p_test, poly)
                    if res: dist_info_test = res
                    break
            
            puntos.append({
                "lat": lat, "lon": lon, "color": "orange", "alerta": False, "pre_alerta": True,
                "sat": "SIMULACRO", "fecha": fecha_sim, "horas": 1,
                "concesion": "Zona de Amortiguamiento", "gtm": convertir_a_gtm(lon, lat),
                "dist_info": dist_info_test,
                "dist_campamento": "N/A"
            })
        elif action_type == "test_monitoreo":
            # Sin puntos, solo fuerza el reporte verde
            pass
            
    # --- LÓGICA NORMAL (DESCARGA NASA) ---
    else:
        # Configurar sesión con reintentos para mayor robustez
        session = requests.Session()
        retry_strategy = Retry(total=3, backoff_factor=2, status_forcelist=[429, 500, 502, 503, 504])
        adapter = HTTPAdapter(max_retries=retry_strategy)
        session.mount("https://", adapter)
        session.mount("http://", adapter)

        satelites = ["MODIS_NRT", "VIIRS_SNPP_NRT", "VIIRS_NOAA20_NRT"]
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
        }
        for sat in satelites:
            try:
                print(f"⬇️ Descargando datos de {sat}...")
                url = f"https://firms.modaps.eosdis.nasa.gov/api/area/csv/{MAP_KEY}/{sat}/-94,13.5,-88,20/3"
                res = session.get(url, headers=headers, timeout=20)
                if res.status_code == 200:
                    lines = res.text.strip().split('\n')[1:]
                    for line in lines:
                        try:
                            # Limpieza robusta de datos
                            d = [x.strip().replace('"', '') for x in line.split(',')]
                            lat, lon = float(d[0]), float(d[1])
                            p = Point(lon, lat)
                            
                            # Verificar ubicación
                            en_paxban = False
                            en_prealerta = False
                            concesion_nombre = "Externa"
                            dist_campamento_str = "N/A"
                            dist_info = None
                            
                            for nom, poly in concesiones.items():
                                if "Paxbán" in nom:
                                    if poly.contains(p):
                                        en_paxban = True
                                        concesion_nombre = "Paxbán"
                                        dist_campamento_str = calcular_campamento_cercano(lon, lat)
                                        break
                                    # Buffer aproximado de 10km (0.09 grados)
                                    elif poly.distance(p) < 0.09:
                                        en_prealerta = True
                                        concesion_nombre = "Zona de Amortiguamiento"
                                        dist_info = calcular_distancia_direccion(p, poly)
                            
                            # Calcular antigüedad
                            dt = datetime.strptime(f"{d[5]} {d[6]}", "%Y-%m-%d %H%M")
                            
                            # Filtro de seguridad: Ignorar fechas futuras (errores de API)
                            if dt > datetime.utcnow(): continue
                            
                            horas = (datetime.utcnow() - dt).total_seconds() / 3600
                            color = "red" if horas <= 24 else "orange" if horas <= 48 else "yellow"
                            
                            # Convertir a hora Guatemala (UTC-6)
                            dt_gt = dt - timedelta(hours=6)
                            fecha_gt = dt_gt.strftime("%d/%m/%Y %H:%M")
                            
                            puntos.append({
                                "lat": lat, "lon": lon, "color": color, "alerta": en_paxban, "pre_alerta": en_prealerta,
                                "sat": sat, "fecha": f"{fecha_gt} (Hora GT)", "horas": horas,
                                "concesion": concesion_nombre,
                                "gtm": convertir_a_gtm(lon, lat),
                                "dist_info": dist_info,
                                "dist_campamento": dist_campamento_str
                            })
                        except: pass
            except requests.exceptions.ConnectionError:
                print(f"⛔ Error de conexión con {sat}. Internet inestable o servidor caído.")
                print("⚠️ Saltando satélites restantes para evitar demoras innecesarias.")
                break
            except Exception as e:
                print(f"⚠️ Error descargando {sat}: {e}", file=sys.stderr)

        if not puntos:
            print("⚠️ Advertencia: No se encontraron datos de incendios en el área seleccionada.")

    # Guardar Metadatos de Actualización (Hora)
    fecha_actual = (datetime.utcnow() - timedelta(hours=6)).strftime("%d/%m/%Y %H:%M")
    with open('metadata.json', 'w') as f:
        json.dump({"last_updated": f"{fecha_actual} (Hora GT)"}, f)

    # Guardar JSON para la web
    with open('incendios.json', 'w') as f: json.dump(puntos, f)
    
    # Solo alertar sobre puntos detectados en las últimas 1.5 horas para evitar duplicados en ejecuciones horarias
    alertas = [p for p in puntos if p['alerta'] and p['horas'] <= 1.5]
    pre_alertas = [p for p in puntos if p.get('pre_alerta') and p['horas'] <= 1.5]
    force_report = os.environ.get("FORCE_REPORT", "false") == "true"
    
    # Generar mapa si es necesario
    img_bytes = None
    if alertas or pre_alertas or force_report:
        img_bytes = generar_mapa_imagen(puntos, concesiones)
        if force_report: guardar_mapa_local(img_bytes)
        if alertas: guardar_bitacora(img_bytes, "alertas", alertas)
        if pre_alertas: guardar_bitacora(img_bytes, "pre_alertas", pre_alertas)

    # --- ENVIAR CORREOS SEGÚN PRIORIDAD ---
    
    # 1. ALERTA ROJA (Incendio DENTRO de Paxbán)
    if alertas:
        msg = f"🔥 <b>ALERTA PAXBAN</b>\n"
        msg += f"<b>Se detectaron {len(alertas)} incendios activos.</b>\n"
        for i, p in enumerate(alertas):
            msg += f"\n🔴 <b>Foco {i+1}:</b>"
            msg += f"\n📍 {p['lat']:.5f}, {p['lon']:.5f}"
            msg += f"\n🗺️ GTM: {p['gtm']}"
            msg += f"\n🛰️ {p['sat']} - {p['fecha']}"
        enviar_alerta_telegram(msg, img_bytes)
        
        html = f"""
        <html>
        <head>
        <style>
            @media print {{
                @page {{ margin: 0.5cm; }}
                body {{ font-family: Arial, sans-serif; font-size: 9pt; }}
                h2 {{ color: #D32F2F; margin-top: 0; font-size: 12pt; margin-bottom: 5px; }}
                .alert-box {{ background-color: #ffcdd2 !important; border-left: 5px solid #D32F2F !important; -webkit-print-color-adjust: exact; padding: 5px !important; margin: 5px 0 !important; }}
                table {{ width: 100%; border-collapse: collapse; font-size: 8pt; margin-bottom: 5px !important; }}
                th {{ background-color: #ef5350 !important; color: white !important; -webkit-print-color-adjust: exact; padding: 2px; }}
                td {{ padding: 2px; border: 1px solid #ddd; }}
                img {{ max-height: 300px !important; width: auto; display: block; margin: 5px auto; }}
                .no-print {{ display: none; }}
            }}
        </style>
        </head>
        <body>
        <div style="font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto;">
            <table style="width: 100%; border-bottom: 2px solid #D32F2F; margin-bottom: 10px;">
                <tr>
                    <td style="width: 80px; padding-bottom: 5px;">
                        <img src="cid:logo_paxban" alt="Logo Paxban" style="width: 70px; height: auto;">
                    </td>
                    <td style="vertical-align: middle; padding-bottom: 5px;">
                        <h2 style="color: #D32F2F; margin: 0; font-size: 18pt;">🔥 ALERTA DE INCENDIO DETECTADO 🔥</h2>
                    </td>
                </tr>
            </table>
            <p style="margin: 5px 0;">Estimado usuario,</p>
            <p style="margin: 5px 0;"><strong>¡Atención!</strong> El sistema Alerta Paxban4 ha identificado <strong>{len(alertas)} foco(s) de incendio activos</strong> dentro de los polígonos de las concesiones monitoreadas.</p>
            <div class="alert-box" style="background-color: #ffcdd2; padding: 10px; border-left: 5px solid #D32F2F; margin: 10px 0;">
                <h3 style="margin: 0; color: #b71c1c; font-size: 14pt;">Resumen de la Alerta</h3>
                <p style="margin: 5px 0 0 0;">Se requiere verificación y acción inmediata.</p>
            </div>
            <h4 style="color: #333; margin: 10px 0 5px 0;">Detalles de los Focos Detectados:</h4>
            <table style="width: 100%; border-collapse: collapse; font-size: 12px;">
                <tr style="background-color: #ef5350; color: white; text-align: left;">
                    <th style="padding: 5px; border: 1px solid #ddd;">#</th><th style="padding: 5px; border: 1px solid #ddd;">Referencia</th><th style="padding: 5px; border: 1px solid #ddd;">Coordenadas</th><th style="padding: 5px; border: 1px solid #ddd;">GTM</th><th style="padding: 5px; border: 1px solid #ddd;">Fecha/Hora</th>
                </tr>"""
        for i, p in enumerate(alertas):
            html += f"""
            <tr style="background-color: {'#ffebee' if i % 2 == 0 else '#ffffff'}; font-size: 11px;">
                <td style="padding: 4px; border: 1px solid #ddd;">{i+1}</td>
                <td style="padding: 4px; border: 1px solid #ddd;"><strong>{p.get('dist_campamento', 'N/A')}</strong></td>
                <td style="padding: 4px; border: 1px solid #ddd;">{p['lat']:.4f}, {p['lon']:.4f}</td>
                <td style="padding: 4px; border: 1px solid #ddd;">{p['gtm']}</td>
                <td style="padding: 4px; border: 1px solid #ddd;">{p['fecha']}</td>
            </tr>"""
        html += "</table>"
        
        # Agregar aviso de pre-alertas si existen simultáneamente
        if pre_alertas:
            html += f"""<div style="background-color: #FFF3E0; padding: 8px; border-left: 4px solid #F57F17; margin: 10px 0;">
                <h4 style="margin: 0; color: #E65100;">⚠️ Adicionalmente: {len(pre_alertas)} focos en Zona de Amortiguamiento</h4>
                <p style="margin: 2px 0 0 0; font-size: 11px;">Se detectó actividad en el perímetro (10km) que requiere vigilancia.</p>
            </div>"""
            
        html += '<p style="margin: 10px 0 5px 0;">A continuación se presenta el mapa de la situación:</p>'
        if img_bytes: html += '<br><img src="cid:mapa_peten" style="max-width: 100%; max-height: 350px; height: auto; border: 1px solid #ddd; border-radius: 5px; display: block; margin: 0 auto;"><br>'
        html += f"""<br><hr style="border: 0; border-top: 1px solid #eee; margin: 10px 0;"><div style="font-size: 11px; color: #666;"><p style="margin: 2px 0;"><b>Sistema de Alerta Temprana Paxban</b><br>Mensaje generado por detección de amenaza.<br>Desarrollado por JR23CR</p><p style="text-align: center; margin-top: 10px;" class="no-print"><a href="https://JR23CR.github.io/alerta-paxban/reportes.html" style="background-color: #D32F2F; color: white; padding: 8px 15px; text-decoration: none; border-radius: 5px; font-weight: bold; font-size: 12px;">📂 Ver Galería de Reportes</a></p></div></div></body></html>"""
        enviar_correo_alerta(html, asunto="🔥 ALERTA DE INCENDIO - Paxban", imagen_mapa=img_bytes)
        
    # 2. PRE-ALERTA AMARILLA (Incendio CERCA de Paxbán)
    elif pre_alertas:
        print(f"📧 Enviando Pre-Alerta por {len(pre_alertas)} focos...")
        msg = f"⚠️ <b>PRE-ALERTA PAXBAN</b>\n"
        msg += f"<b>Actividad en zona de amortiguamiento ({len(pre_alertas)} focos).</b>\n"
        for i, p in enumerate(pre_alertas):
            msg += f"\n🟠 <b>Foco {i+1}:</b>"
            if p.get('dist_info'):
                msg += f"\n📏 {p['dist_info']}"
            msg += f"\n📍 {p['lat']:.5f}, {p['lon']:.5f}"
            msg += f"\n🗺️ GTM: {p['gtm']}"
            msg += f"\n🛰️ {p['sat']} - {p['fecha']}"
        enviar_alerta_telegram(msg, img_bytes)
        
        html = f"""
        <html>
        <head>
        <style>
            @media print {{
                @page {{ margin: 0.5cm; }}
                body {{ font-family: Arial, sans-serif; font-size: 9pt; }}
                h2 {{ color: #F57F17; margin-top: 0; font-size: 12pt; margin-bottom: 5px; }}
                .alert-box {{ background-color: #FFF3E0 !important; border-left: 5px solid #F57F17 !important; -webkit-print-color-adjust: exact; padding: 5px !important; margin: 5px 0 !important; }}
                table {{ width: 100%; border-collapse: collapse; font-size: 8pt; margin-bottom: 5px !important; }}
                th {{ background-color: #FFB74D !important; color: white !important; -webkit-print-color-adjust: exact; padding: 2px; }}
                td {{ padding: 2px; border: 1px solid #ddd; }}
                img {{ max-height: 300px !important; width: auto; display: block; margin: 5px auto; }}
                .no-print {{ display: none; }}
            }}
        </style>
        </head>
        <body>
        <div style="font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto;">
            <table style="width: 100%; border-bottom: 2px solid #F57F17; margin-bottom: 10px;">
                <tr>
                    <td style="width: 80px; padding-bottom: 5px;">
                        <img src="cid:logo_paxban" alt="Logo Paxban" style="width: 70px; height: auto;">
                    </td>
                    <td style="vertical-align: middle; padding-bottom: 5px;">
                        <h2 style="color: #F57F17; margin: 0; font-size: 18pt;">⚠️ PRE-ALERTA DE INCENDIO</h2>
                    </td>
                </tr>
            </table>
            <p style="margin: 5px 0;">Estimado usuario,</p>
            <p style="margin: 5px 0;">El sistema ha detectado <strong>{len(pre_alertas)} foco(s) de calor</strong> en la zona de amortiguamiento (10 km) de la concesión.</p>
            <div class="alert-box" style="background-color: #FFF3E0; padding: 10px; border-left: 5px solid #F57F17; margin: 10px 0;">
                <h3 style="margin: 0; color: #E65100; font-size: 14pt;">Zona de Vigilancia</h3>
                <p style="margin: 5px 0 0 0;">Se recomienda monitorear la evolución de estos puntos.</p>
            </div>
            <h4 style="color: #333; margin: 10px 0 5px 0;">Detalles:</h4>
            <table style="width: 100%; border-collapse: collapse; font-size: 12px;">
                <tr style="background-color: #FFB74D; color: white; text-align: left;">
                    <th style="padding: 5px;">#</th><th style="padding: 5px;">Ubicación</th><th style="padding: 5px;">Coordenadas</th><th style="padding: 5px;">GTM</th><th style="padding: 5px;">Fecha/Hora</th>
                </tr>"""
        for i, p in enumerate(pre_alertas):
            dist_txt = p.get('dist_info', 'Zona de Amortiguamiento')
            html += f"""<tr style="background-color: #ffffff; font-size: 11px;"><td style="padding: 4px; border: 1px solid #ddd;">{i+1}</td><td style="padding: 4px; border: 1px solid #ddd;">{dist_txt}</td><td style="padding: 4px; border: 1px solid #ddd;">{p['lat']:.4f}, {p['lon']:.4f}</td><td style="padding: 4px; border: 1px solid #ddd;">{p['gtm']}</td><td style="padding: 4px; border: 1px solid #ddd;">{p['fecha']}</td></tr>"""
        html += "</table>"
        if img_bytes: html += '<br><img src="cid:mapa_peten" style="max-width: 100%; max-height: 350px; height: auto; border: 1px solid #ddd; border-radius: 5px; display: block; margin: 0 auto;"><br>'
        html += f"""<br><hr style="border: 0; border-top: 1px solid #eee; margin: 10px 0;"><div style="font-size: 11px; color: #666;"><p style="margin: 2px 0;"><b>Sistema de Alerta Temprana Paxban</b><br>Mensaje de advertencia preventiva.<br>Desarrollado por JR23CR</p><p style="text-align: center; margin-top: 10px;" class="no-print"><a href="https://JR23CR.github.io/alerta-paxban/reportes.html" style="background-color: #F57F17; color: white; padding: 8px 15px; text-decoration: none; border-radius: 5px; font-weight: bold; font-size: 12px;">📂 Ver Galería de Reportes</a></p></div></div></body></html>"""
        enviar_correo_alerta(html, asunto="⚠️ PRE-ALERTA DE INCENDIO - Paxban", imagen_mapa=img_bytes)

    # 3. REPORTE DIARIO VERDE (Sin amenazas o solo monitoreo)
    elif force_report:
        fecha_hora = (datetime.utcnow() - timedelta(hours=6)).strftime("%d/%m/%Y %H:%M")
        razon = os.environ.get("REPORT_REASON", "automáticamente")
        
        # Solo enviar notificaciones si NO es el reporte automático programado de las 4 PM
        # Y si NO es una generación de reporte mensual (para evitar doble correo)
        es_reporte_programado = "automáticamente" in razon and "Reporte Diario" in razon
        es_generacion_mensual = action_type == "reporte_mensual"
        
        if not es_reporte_programado and not es_generacion_mensual:
            # Mensaje Telegram
            msg = f"🛰️ <b>Reporte de Monitoreo Satelital</b>\n\n" \
                  f"✅ <b>Estado: Sin Amenazas Detectadas</b>\n" \
                  f"No se han identificado focos activos dentro de las concesiones.\n\n" \
                  f"📍 Puntos analizados: {len(puntos)}\n" \
                  f"🕒 Hora: {fecha_hora}\n\n" \
                  f"Sistema de Alerta Temprana Paxban"
            enviar_alerta_telegram(msg, img_bytes)
        
        # HTML Correo (Tu diseño)
        html = f"""
        <html>
        <head>
        <style>
            @media print {{
                @page {{ margin: 0.5cm; }}
                body {{ font-family: Arial, sans-serif; font-size: 9pt; }}
                h2 {{ color: #2E7D32; margin-top: 0; font-size: 12pt; margin-bottom: 5px; }}
                .status-box {{ background-color: #e8f5e9 !important; border-left: 5px solid #2e7d32 !important; -webkit-print-color-adjust: exact; padding: 5px !important; margin: 5px 0 !important; }}
                img {{ max-height: 320px !important; width: auto; display: block; margin: 5px auto; }}
                p {{ margin: 2px 0; }}
                table {{ margin-bottom: 5px !important; }}
                td {{ padding-bottom: 0 !important; }}
                .no-print {{ display: none; }}
            }}
        </style>
        </head>
        <body>
        <div style="font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto;">
            <table style="width: 100%; border-bottom: 2px solid #2E7D32; margin-bottom: 10px;">
                <tr>
                    <td style="width: 80px; padding-bottom: 5px;">
                        <img src="cid:logo_paxban" alt="Logo Paxban" style="width: 70px; height: auto;">
                    </td>
                    <td style="vertical-align: middle; padding-bottom: 5px;">
                         <h2 style="color: #2E7D32; margin: 0; font-size: 18pt;">Reporte de Monitoreo Satelital</h2>
                    </td>
                </tr>
            </table>
            <p style="margin: 5px 0;">Estimado usuario,</p>
            <p style="margin: 5px 0;">El sistema Alerta Paxban ha completado el análisis de los datos satelitales más recientes.</p>
            
            <div class="status-box" style="background-color: #e8f5e9; padding: 10px; border-left: 5px solid #2e7d32; margin: 15px 0;">
                <h3 style="margin: 0; color: #1b5e20; font-size: 14pt;">✅ Estado: Sin Amenazas Detectadas</h3>
                <p style="margin: 5px 0 0 0;">No se han identificado focos de incendio activos dentro de los polígonos de las concesiones forestales monitoreadas.</p>
            </div>

            <p style="margin: 5px 0;">
                <b>Puntos analizados en la región:</b> {len(puntos)}<br>
                <b>Hora del reporte:</b> {fecha_hora}
            </p>

            <p style="margin: 5px 0;">A continuación, se presenta el Mapa de Situación Actual en Petén, mostrando la actividad térmica general en la región. Los colores indican la antigüedad del punto de calor (Rojo: &lt;24h, Naranja: &lt;48h, Amarillo: &lt;72h).</p>
        """
        
        if img_bytes: html += '<br><img src="cid:mapa_peten" style="max-width: 100%; max-height: 350px; height: auto; border: 1px solid #ddd; border-radius: 5px; display: block; margin: 0 auto;"><br>'
        
        html += f"""
            <br>
            <hr style="border: 0; border-top: 1px solid #eee; margin: 10px 0;">
            <div style="font-size: 11px; color: #666;">
                <p style="margin: 2px 0;">
                    <b>Sistema de Alerta Temprana Paxban</b><br>
                    Mensaje generado {razon}.<br>
                    Desarrollado por JR23CR
                </p>
                <p style="text-align: center; margin-top: 10px;" class="no-print">
                    <a href="https://JR23CR.github.io/alerta-paxban/reportes.html" style="background-color: #2E7D32; color: white; padding: 8px 15px; text-decoration: none; border-radius: 5px; font-weight: bold; font-size: 12px;">📂 Ver Galería de Reportes</a>
                </p>
            </div>
        </div>
        </body>
        </html>
        """
        if not es_reporte_programado and not es_generacion_mensual:
            enviar_correo_alerta(html, asunto="Reporte de Monitoreo Satelital", imagen_mapa=img_bytes)
        else:
            print("ℹ️ Reporte diario procesado (Mapa guardado), pero se omite el envío de correo (Programado o Mensual).")

    # Reporte Mensual si se solicita
    if os.environ.get("ACTION_TYPE") == "reporte_mensual":
        generar_reporte_mensual(concesiones)

    # Generar Galería SIEMPRE (Al final para incluir el reporte mensual si se generó)
    limpiar_descargas_antiguas()
    time.sleep(1) # Esperar un segundo para liberar bloqueos de archivos
    generar_galeria_html()

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print("❌ ERROR FATAL EN EL SCRIPT:", file=sys.stderr)
        traceback.print_exc()
        # No salimos con error para permitir que GitHub Pages intente desplegar lo que haya
        sys.exit(0) 
